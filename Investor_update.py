import os
import streamlit as st
from datetime import datetime
import tempfile
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openai
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from dotenv import load_dotenv

# Try to import PDF reader
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Load environment variables
env_path = os.path.join(os.path.dirname(__file__), ".env")
load_dotenv(dotenv_path=env_path)


def load_documents_from_uploads(uploaded_files, storage_key="investor"):
    """Load documents from uploaded files and store in session state with app-specific key"""
    docs = []
    
    if not uploaded_files:
        return docs
    
    # Use app-specific storage keys
    docs_key = f'processed_docs_{storage_key}'
    names_key = f'processed_file_names_{storage_key}'
    
    # Check if files are already processed in session state
    current_file_names = [f.name for f in uploaded_files]
    if docs_key in st.session_state and names_key in st.session_state:
        if st.session_state[names_key] == current_file_names:
            return st.session_state[docs_key]
    
    for uploaded_file in uploaded_files:
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name
            
            # Process based on file type
            if uploaded_file.name.lower().endswith('.pdf'):
                if PDF_AVAILABLE:
                    result = load_pdf_safe(tmp_path)
                    docs.extend(result)
                else:
                    with st.sidebar:
                        st.warning(f"⏭️ Skipping PDF {uploaded_file.name} (PyPDF2 not installed)")
            
            elif uploaded_file.name.lower().endswith('.docx'):
                result = load_docx_safe(tmp_path)
                docs.extend(result)
            
            elif uploaded_file.name.lower().endswith(('.txt', '.md')):
                result = load_text_safe(tmp_path)
                docs.extend(result)
            
            # Clean up temp file
            os.unlink(tmp_path)
            
        except Exception as e:
            with st.sidebar:
                st.warning(f"❌ Error loading {uploaded_file.name}: {str(e)}")
    
    # Store processed documents in session state with app-specific keys
    st.session_state[docs_key] = docs
    st.session_state[names_key] = current_file_names
    
    return docs


def load_pdf_safe(file_path):
    """Safely load PDF using PyPDF2"""
    docs = []
    try:
        reader = PdfReader(str(file_path))
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e:
                with st.sidebar:
                    st.warning(f"Error reading page {page_num + 1} of {Path(file_path).name}: {e}")

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "pdf"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading PDF {Path(file_path).name}: {e}")
    return docs


def load_docx_safe(file_path):
    """Safely load DOCX file"""
    docs = []
    try:
        doc = Document(str(file_path))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "docx"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading DOCX {Path(file_path).name}: {e}")
    return docs


def load_text_safe(file_path):
    """Safely load text file"""
    docs = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "text"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading text file {Path(file_path).name}: {e}")
    return docs


@st.cache_resource
def load_rag_engine_with_docs(_docs, temperature):
    """Initialize RAG engine with provided documents"""
    if not _docs:
        # Create a simple LLM without retrieval
        return ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
        )

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(_docs)

    try:
        # Create embeddings + vector store
        embeddings = OpenAIEmbeddings(api_key=os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY"))
        vectorstore = FAISS.from_documents(chunks, embeddings)

        # Setup LLM and RetrievalQA
        llm = ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
        )

        rag_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type='stuff',
            retriever=vectorstore.as_retriever()
        )

        return rag_chain

    except Exception as e:
        with st.sidebar:
            st.error(f"Error creating RAG engine: {e}")
        # Fallback to simple LLM
        return ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            api_key=os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")
        )


def render_investor_ui():
    # OpenAI key
    openai.api_key = os.getenv("OPENAI_API_KEY") or st.secrets.get("OPENAI_API_KEY")

    # Constants
    UPDATE_TYPES = ["Monthly Update", "Quarterly Update", "Milestone Update", "Board Update"]
    LENGTH_MAP = {"Brief": "200-300 words", "Standard": "400-600 words", "Detailed": "700-900 words"}

    # Comprehensive system prompt for Naware
    NAWARE_SYSTEM_PROMPT = """
    You are generating professional investor updates for Naware, a Minneapolis-based deep tech startup founded by Mark Boysen in May 2024.

    COMPANY CONTEXT:
    - Mission: Revolutionize weed control by eliminating harmful chemicals using AI, robotics, and steam
    - Target Market: B2B clients including golf courses, municipalities, and commercial lawn care companies
    - Technology: AI-powered weed detection system (85-90% accuracy) with precision steam delivery
    - Market Opportunity: $34B global weed control market, with focus on $5B lawn care segment
    - Business Model: Hardware units priced at $28K+ with 46%+ gross margins, targeting 100 units by Q1 2026
    - Leadership Team: Mark Boysen (Founder/CEO), Sudee (Robotics Lead), Santosh (AI Specialist), Kelsey (Sales), Obaid (Technical)

    PROFESSIONAL WRITING STYLE REQUIREMENTS:
    - Formal, professional tone appropriate for institutional investors
    - Factual, data-driven reporting with specific metrics
    - Clear section structure with executive summary approach
    - Third-person perspective with occasional first-person for leadership voice
    - Conservative language that builds confidence without overpromising
    - Include quantitative results, timelines, and measurable outcomes
    - Professional business terminology and industry-standard KPIs
    - Balanced reporting that acknowledges both achievements and challenges

    REALISTIC BUSINESS METRICS:
    - Lead Generation: 75-200 qualified prospects
    - Sales Pipeline: 8-25 active demonstrations per quarter
    - Revenue Range: $0-$200K for current development stage
    - Manufacturing Capacity: 25-1000 units depending on scale phase
    - Customer Pipeline: 10-25 qualified prospects in active evaluation
    - Team Size: 5-10 professionals across technical and commercial functions

    Write from the CEO's perspective providing transparent, professional updates to the investment community.
    """

    # --- Main UI ---
    st.title("📈 Naware Professional Investor Updates")
    st.markdown("Generate institutional-grade investor communications with comprehensive business metrics and professional formatting.")

    # --- Sidebar Configuration ---
    st.sidebar.header("⚙️ Configuration")
    temperature = st.sidebar.slider("Temperature", 0.1, 1.0, 0.7, 0.1)
    update_length = st.sidebar.select_slider("Update Length", ["Brief", "Standard", "Detailed"], "Standard")
    update_type = st.sidebar.selectbox("Update Type", UPDATE_TYPES, index=0)
    company_name = st.sidebar.text_input("Company Name", "Naware")
    update_date = st.sidebar.date_input("Update Date", datetime.now())

    # Advanced options
    with st.sidebar.expander("🔧 Advanced Options"):
        include_metrics = st.checkbox("Include Metrics Dashboard", True)
        include_financials = st.checkbox("Include Financial Summary", True)
        tone = st.selectbox("Tone", ["Optimistic", "Balanced", "Conservative"], index=0)

    # --- Document Upload in Sidebar ---
    st.sidebar.header("📁 Upload Documents")
    uploaded_files = st.sidebar.file_uploader(
        "Company docs (PDF, DOCX, TXT, MD)",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'txt', 'md'],
        help="Upload company documents for investor update context",
        key="investor_upload"
    )
    
    # Show currently stored files for Investor Updates
    investor_docs_key = 'processed_file_names_investor'
    if investor_docs_key in st.session_state and st.session_state[investor_docs_key]:
        st.sidebar.subheader("📄 Investor Files")
        for file_name in st.session_state[investor_docs_key]:
            st.sidebar.write(f"✅ {file_name}")
        
        # Add clear button for investor files
        if st.sidebar.button("🗑️ Clear Investor Files"):
            if 'processed_docs_investor' in st.session_state:
                del st.session_state['processed_docs_investor']
            if 'processed_file_names_investor' in st.session_state:
                del st.session_state['processed_file_names_investor']
            st.rerun()
    
    # Document processing - check session state first
    all_docs = []
    if uploaded_files:
        all_docs = load_documents_from_uploads(uploaded_files, "investor")
    elif 'processed_docs_investor' in st.session_state:
        # Use previously processed documents if no new files uploaded
        all_docs = st.session_state['processed_docs_investor']

    # Initialize RAG engine with uploaded documents
    try:
        rag_chain = load_rag_engine_with_docs(all_docs, temperature)
    except Exception as e:
        st.sidebar.error(f"❌ RAG Engine Error: {e}")
        rag_chain = None

    # Professional templates for investor communications
    st.header("📋 Professional Update Templates")
    col1, col2, col3, col4 = st.columns(4)

    with col1:
        if st.button("📊 Quarterly Review"):
            st.session_state['topics'] = ["Executive Summary", "Financial Performance", "Operational Highlights", "Strategic Objectives", "Risk Assessment"]

    with col2:
        if st.button("🎯 Progress Report"):
            st.session_state['topics'] = ["Milestone Achievements", "Product Development Status", "Commercial Pipeline", "Operational Metrics"]

    with col3:
        if st.button("💼 Board Update"):
            st.session_state['topics'] = ["Strategic Overview", "Financial Summary", "Team & Operations", "Market Position", "Forward Guidance"]

    with col4:
        if st.button("📈 Performance Review"):
            st.session_state['topics'] = ["KPI Dashboard", "Revenue Analysis", "Customer Acquisition", "Technology Progress", "Investment Utilization"]

    # Topic management
    if 'topics' not in st.session_state:
        st.session_state['topics'] = []

    st.header("📝 Update Topics")
    with st.form(key='topic_form', clear_on_submit=True):
        col1, col2 = st.columns([3, 1])
        with col1:
            new_topic = st.text_input("Enter a topic to cover in the update")
        with col2:
            st.write("")
            st.write("")
            if st.form_submit_button("➕ Add Topic") and new_topic.strip():
                st.session_state['topics'].append(new_topic.strip())

    # Display and edit topics
    if st.session_state['topics']:
        st.subheader("Current Topics:")
        for idx, topic in enumerate(st.session_state['topics']):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.session_state['topics'][idx] = st.text_input(
                    f"Topic {idx + 1}", value=topic, key=f"topic_{idx}"
                ).strip()
            with col2:
                st.write("")
                st.write("")
                if st.button("🗑️", key=f"delete_{idx}"):
                    st.session_state['topics'].pop(idx)
                    st.rerun()

    # Generate function
    def generate_investor_update():
        """Generate the investor update content"""
        if not st.session_state['topics']:
            st.error("Please add at least one topic.")
            return None, None

        if not rag_chain:
            st.error("RAG engine not initialized. Please check your OpenAI API key.")
            return None, None

        all_sections = []
        progress_bar = st.progress(0)
        status_text = st.empty()

        for idx, topic in enumerate(st.session_state['topics']):
            status_text.text(f"Generating content for: {topic}")

            # Create professional prompt for each topic
            prompt = f"""
            {NAWARE_SYSTEM_PROMPT}

            Generate a {LENGTH_MAP[update_length]} professional section for an investor update covering "{topic}".

            Requirements:
            - Use formal business language appropriate for institutional investors
            - Include specific quantitative metrics and performance data
            - Structure with clear headings and bullet points where appropriate
            - Focus on measurable outcomes, timelines, and business impact
            - Provide context for achievements within market conditions
            - Address both progress and challenges transparently
            - Use professional terminology and avoid casual expressions
            - Include forward-looking statements with appropriate caveats

            Section Topic: {topic}
            Update Type: {update_type}
            Communication Tone: {tone} and Professional
            Reporting Period: {update_date.strftime('%B %Y')}

            Format the response with clear structure and professional business language suitable for investor communications.
            """

            try:
                if hasattr(rag_chain, 'invoke'):
                    result = rag_chain.invoke({"query": prompt})
                elif hasattr(rag_chain, 'run'):
                    result = rag_chain.run(prompt)
                else:
                    # Fallback for simple ChatOpenAI
                    result = rag_chain.predict(prompt)

                # Extract text content from result
                if isinstance(result, dict):
                    result_text = result.get('result', '') or result.get('answer', '') or str(result)
                else:
                    result_text = str(result)

                # Clean and structure the result
                paragraphs = [p.strip() for p in re.split(r'\n\n|\n', result_text) if p.strip()]
                all_sections.append((topic, paragraphs))

            except Exception as e:
                st.error(f"Error generating content for '{topic}': {e}")
                continue

            progress_bar.progress((idx + 1) / len(st.session_state['topics']))

        status_text.text("Update generated successfully!")
        return all_sections, create_docx_update(all_sections)

    def create_docx_update(sections):
        """Create a professionally formatted DOCX document for investors"""
        doc = Document()

        # Professional header
        header = doc.add_heading(f"{company_name} Investor Update", 0)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Date and type
        date_para = doc.add_paragraph(f"{update_type}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para = doc.add_paragraph(f"{update_date.strftime('%B %d, %Y')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Space

        # Professional opening
        doc.add_paragraph("Dear Investors,")
        doc.add_paragraph()
        opening = doc.add_paragraph(
            f"I am pleased to provide you with our {update_type.lower()} covering key developments "
            f"across our operational and strategic initiatives. This update reflects our continued "
            f"progress toward our stated objectives and market milestones."
        )
        doc.add_paragraph()

        # Executive Summary if multiple sections
        if len(sections) > 2:
            doc.add_heading("Executive Summary", level=1)
            summary_para = doc.add_paragraph(
                "This update highlights significant progress in product development, commercial "
                "pipeline advancement, and operational scaling. Key achievements include enhanced "
                "technology performance, expanded customer engagement, and strengthened market position."
            )
            doc.add_paragraph()

        # Content sections with professional formatting
        for topic, paragraphs in sections:
            doc.add_heading(topic, level=1)
            for para in paragraphs:
                if para and len(para.strip()) > 10:  # Only substantial paragraphs
                    formatted_para = doc.add_paragraph(para)
            doc.add_paragraph()

        # Professional closing section
        doc.add_heading("Looking Forward", level=1)
        forward_para = doc.add_paragraph(
            "We remain focused on executing our strategic roadmap and delivering measurable "
            "value to our stakeholders. Our team continues to advance our technology platform "
            "while building sustainable commercial relationships that will drive long-term growth."
        )
        doc.add_paragraph()

        # Contact and availability
        doc.add_paragraph(
            "As always, I welcome the opportunity to discuss our progress in greater detail. "
            "Please feel free to reach out with any questions or to schedule a call."
        )
        doc.add_paragraph()

        # Professional signature
        doc.add_paragraph("Respectfully,")
        doc.add_paragraph()
        doc.add_paragraph("Mark Boysen")
        doc.add_paragraph("Chief Executive Officer")
        doc.add_paragraph(f"{company_name}")

        # Save to temp file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            return tmp.name

    # Generate button and results
    if st.button("🚀 Generate Investor Update", type="primary"):
        with st.spinner("Generating your investor update..."):
            sections, docx_path = generate_investor_update()

            if sections and docx_path:
                # Download button
                with open(docx_path, 'rb') as file:
                    st.download_button(
                        label="📥 Download as DOCX",
                        data=file.read(),
                        file_name=f"{company_name.replace(' ', '_')}_Investor_Update_{update_date.strftime('%Y-%m-%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                st.success("✅ Investor update generated successfully!")

                # Professional preview
                st.markdown("### 📋 Document Preview")

                # Header preview
                st.markdown(f"<h2 style='text-align: center'>{company_name} Investor Update</h2>", unsafe_allow_html=True)
                st.markdown(f"<p style='text-align: center'><strong>{update_type}</strong></p>", unsafe_allow_html=True)
                st.markdown(f"<p style='text-align: center'>{update_date.strftime('%B %d, %Y')}</p>", unsafe_allow_html=True)
                st.markdown("---")

                # Professional opening
                st.markdown("**Dear Investors,**")
                st.markdown(f"I am pleased to provide you with our {update_type.lower()} covering key developments across our operational and strategic initiatives.")
                st.markdown("")

                # Content sections
                for topic, paragraphs in sections:
                    st.markdown(f"### {topic}")
                    for para in paragraphs:
                        if para and len(para.strip()) > 10:
                            st.markdown(para)
                    st.markdown("")

                # Professional closing preview
                st.markdown("### Looking Forward")
                st.markdown("We remain focused on executing our strategic roadmap and delivering measurable value to our stakeholders.")
                st.markdown("")
                st.markdown("Respectfully,")
                st.markdown("**Mark Boysen**  \nChief Executive Officer  \nNaware")
                st.markdown("---")

                # Clear topics after successful generation
                if st.button("🗑️ Clear Topics"):
                    st.session_state['topics'] = []
                    st.rerun()

    # Footer
    st.markdown("---")
    st.markdown("*Naware Professional Investor Communications*")
