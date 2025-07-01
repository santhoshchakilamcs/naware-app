import os
import streamlit as st
from datetime import datetime
import tempfile
import re
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
# import openai
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from dotenv import load_dotenv

# Try to import PDF reader
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

load_dotenv()


def load_documents_from_uploads(uploaded_files, storage_key="newsletter"):
    """Load documents from uploaded files and store in session state with app-specific key"""
    docs = []
    
    if not uploaded_files:
        return docs
    
    # Use app-specific storage keys
    docs_key = f'processed_docs_{storage_key}'
    names_key = f'processed_file_names_{storage_key}'
    
    # Check if files are already processed in session state
    current_file_names = [f.name for f in uploaded_files]
    if docs_key in st.session_state and names_key in st.session_state:
        if st.session_state[names_key] == current_file_names:
            return st.session_state[docs_key]
    
    for uploaded_file in uploaded_files:
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp:
                tmp.write(uploaded_file.read())
                tmp_path = tmp.name
            
            # Process based on file type
            if uploaded_file.name.lower().endswith('.pdf'):
                if PDF_AVAILABLE:
                    result = load_pdf_safe(tmp_path)
                    docs.extend(result)
                else:
                    with st.sidebar:
                        st.warning(f"⏭️ Skipping PDF {uploaded_file.name} (PyPDF2 not installed)")
            
            elif uploaded_file.name.lower().endswith('.docx'):
                result = load_docx_safe(tmp_path)
                docs.extend(result)
            
            elif uploaded_file.name.lower().endswith(('.txt', '.md')):
                result = load_text_safe(tmp_path)
                docs.extend(result)
            
            # Clean up temp file
            os.unlink(tmp_path)
            
        except Exception as e:
            with st.sidebar:
                st.warning(f"❌ Error loading {uploaded_file.name}: {str(e)}")
    
    # Store processed documents in session state with app-specific keys
    st.session_state[docs_key] = docs
    st.session_state[names_key] = current_file_names
    
    return docs


def load_pdf_safe(file_path):
    """Safely load PDF using PyPDF2"""
    docs = []
    try:
        reader = PdfReader(str(file_path))
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e:
                with st.sidebar:
                    st.warning(f"Error reading page {page_num + 1} of {Path(file_path).name}: {e}")

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "pdf"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading PDF {Path(file_path).name}: {e}")
    return docs


def load_docx_safe(file_path):
    """Safely load DOCX file"""
    docs = []
    try:
        doc = Document(str(file_path))
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "docx"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading DOCX {Path(file_path).name}: {e}")
    return docs


def load_text_safe(file_path):
    """Safely load text file"""
    docs = []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        if text.strip():
            docs.append(LangchainDocument(
                page_content=text,
                metadata={"source": str(file_path), "type": "text"}
            ))
    except Exception as e:
        with st.sidebar:
            st.error(f"Error reading text file {Path(file_path).name}: {e}")
    return docs


@st.cache_resource
def load_rag_engine_with_docs(_docs, temperature):
    """Load RAG engine with provided documents"""
    if not _docs:
        # Create a simple LLM without retrieval
        return ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            openai_api_key=st.secrets.get("OPENAI_API_KEY")
        )

    # Split into chunks
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_documents(_docs)

    try:
        # Create embeddings + vector store
        embeddings = OpenAIEmbeddings(openai_api_key=st.secrets.get("OPENAI_API_KEY"))
        vectorstore = FAISS.from_documents(chunks, embeddings)

        # Setup LLM and RetrievalQA
        llm = ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            openai_api_key=st.secrets.get("OPENAI_API_KEY")
        )

        rag_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type='stuff',
            retriever=vectorstore.as_retriever()
        )

        return rag_chain

    except Exception as e:
        with st.sidebar:
            st.error(f"Error creating RAG engine: {e}")
        # Fallback to simple LLM
        return ChatOpenAI(
            model_name='gpt-4o-mini',
            temperature=temperature,
            openai_api_key=st.secrets.get("OPENAI_API_KEY")
        )


def render_newsletter_ui():
    # Constants - COMPANY_DOC is now used for file upload configuration
    COMPANY_DOC = st.secrets.get("COMPANY_DOC", "Upload your company documents for newsletter context")
    LENGTH_MAP = {"Short": "150-200 words", "Medium": "300-400 words", "Long": "500-600 words"}
    STYLE_EXAMPLE = (
        "Roller-Coaster Highlights from the Week:\n"
        ":tools: The Wipe-All Beast: This version is all about brute force...\n"
        ":robot: Gen6's Glow-Up: Gen6 has been busy...\n"
        ":chipmunk: Field Testing Adventures: Let's just say..."
    )

    # --- Main UI ---
    st.title("📰 Naware Newsletter Generator")
    st.markdown("Generate context-rich, humorous newsletters using your company docs and OpenAI.")

    # Sidebar settings
    st.sidebar.header("Configuration")
    temperature = st.sidebar.slider("Temperature", 0.1, 1.0, 0.7, 0.1)
    article_length = st.sidebar.select_slider("Article Length", ["Short", "Medium", "Long"], "Medium")
    company_name = st.sidebar.text_input("Company Name", "Naware")
    newsletter_date = st.sidebar.date_input("Newsletter Date", datetime.now())

    # --- Document Upload in Sidebar ---
    st.sidebar.header("📁 Upload Documents")
    uploaded_files = st.sidebar.file_uploader(
        "Company docs (PDF, DOCX, TXT, MD)",
        accept_multiple_files=True,
        type=['pdf', 'docx', 'txt', 'md'],
        help="Upload company documents for newsletter context",
        key="newsletter_upload"
    )
    
    # Show currently stored files for Newsletter
    newsletter_docs_key = 'processed_file_names_newsletter'
    if newsletter_docs_key in st.session_state and st.session_state[newsletter_docs_key]:
        st.sidebar.subheader("📄 Newsletter Files")
        for file_name in st.session_state[newsletter_docs_key]:
            st.sidebar.write(f"✅ {file_name}")
        
        # Add clear button for newsletter files
        if st.sidebar.button("🗑️ Clear Newsletter Files"):
            if 'processed_docs_newsletter' in st.session_state:
                del st.session_state['processed_docs_newsletter']
            if 'processed_file_names_newsletter' in st.session_state:
                del st.session_state['processed_file_names_newsletter']
            st.rerun()
    
    # Document processing - check session state first
    all_docs = []
    if uploaded_files:
        all_docs = load_documents_from_uploads(uploaded_files, "newsletter")
    elif 'processed_docs_newsletter' in st.session_state:
        # Use previously processed documents if no new files uploaded
        all_docs = st.session_state['processed_docs_newsletter']

    # Initialize RAG engine with uploaded documents
    rag_chain = load_rag_engine_with_docs(all_docs, temperature)

    # Topic entry management - Newsletter specific
    if 'newsletter_topics' not in st.session_state:
        st.session_state['newsletter_topics'] = []

    st.header("Newsletter Topics")
    with st.form(key='newsletter_topic_form', clear_on_submit=True):
        new_topic = st.text_input("Enter a topic")
        if st.form_submit_button("Add Topic") and new_topic.strip():
            st.session_state['newsletter_topics'].append(new_topic.strip())

    # Display current topics with edit capability
    if st.session_state['newsletter_topics']:
        st.subheader("Current Topics:")
        for idx, topic in enumerate(st.session_state['newsletter_topics']):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.session_state['newsletter_topics'][idx] = st.text_input(f"Topic {idx + 1}", value=topic, key=f"newsletter_topic_{idx}").strip()
            with col2:
                st.write("")
                st.write("")
                if st.button("🗑️", key=f"newsletter_delete_{idx}"):
                    st.session_state['newsletter_topics'].pop(idx)
                    st.rerun()

    # Generate action
    def on_generate():
        if not st.session_state['newsletter_topics']:
            st.error("Please add at least one topic.")
            return

        st.info("Generating newsletter content...")
        all_topics = []

        for topic in st.session_state['newsletter_topics']:
            prompt = (
                f"You are a newsletter writer for {company_name}. "
                f"Write a {LENGTH_MAP[article_length]} newsletter article about '{topic}' in a lighthearted, humorous tone, "
                f"using creative subheadings and emojis. Follow this style example:\n{STYLE_EXAMPLE}"
            )

            try:
                if hasattr(rag_chain, 'invoke'):
                    result = rag_chain.invoke(prompt)  # Pass string directly, not dict
                elif hasattr(rag_chain, 'run'):
                    result = rag_chain.run(prompt)
                else:
                    # Fallback for simple ChatOpenAI
                    result = rag_chain.predict(prompt)

                # Extract text content from result
                if isinstance(result, dict):
                    result_text = result.get('result', '') or result.get('answer', '') or str(result)
                elif hasattr(result, 'content'):
                    # Handle LangChain response object with content attribute
                    result_text = result.content
                else:
                    result_text = str(result)

                paras = [p.strip() for p in re.split(r'\n\n|\n', result_text) if p.strip()]
                all_topics.append((topic, paras))

            except Exception as e:
                st.error(f"Error generating content for '{topic}': {e}")
                continue

        # Generate DOCX
        doc = Document()
        doc.styles['Normal'].font.name = 'Arial'
        doc.add_heading(f"{company_name} Newsletter", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(newsletter_date.strftime('%B %d, %Y')).alignment = WD_ALIGN_PARAGRAPH.CENTER

        for topic, paras in all_topics:
            doc.add_heading(topic, level=2)
            for para in paras:
                if para:  # Only add non-empty paragraphs
                    doc.add_paragraph(para).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph("That's a Wrap (for Now)").alignment = WD_ALIGN_PARAGRAPH.CENTER

        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            with open(tmp.name, 'rb') as file:
                file_data = file.read()
                
        # Provide download button
        download_clicked = st.download_button(
            "📥 Download as DOCX",
            data=file_data,
            file_name=f"{company_name.replace(' ', '_')}_Newsletter_{newsletter_date.strftime('%Y-%m-%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Clear topics after download
        if download_clicked:
            st.session_state['newsletter_topics'] = []

        st.success("Newsletter generated! Download your DOCX above.")
        st.markdown("### Preview")
        for topic, paras in all_topics:
            st.subheader(topic)
            for para in paras:
                if para:
                    st.write(para)
            st.write("---")

        # Clear topics after successful generation
        st.session_state['newsletter_topics'] = []

    if st.button("Generate Newsletter", type="primary"):
        on_generate()
